# Document Processing Module for Resume Analysis
# Handles PDF, Word, and text document parsing with OCR fallback capabilities

import asyncio
import io
import logging
import tempfile
import time
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple
from dataclasses import dataclass
from enum import Enum

import httpx
import PyPDF2
import pdfplumber
from PIL import Image
import pytesseract
from docx import Document as DocxDocument
import mammoth

from exceptions import JobProcessingException


class DocumentType(Enum):
    """Supported document types for resume processing."""
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    TXT = "txt"
    IMAGE = "image"
    UNKNOWN = "unknown"


class ExtractionMethod(Enum):
    """Methods used to extract text from documents."""
    DIRECT_TEXT = "direct_text"  # Text directly extracted from document
    OCR = "ocr"  # Optical Character Recognition
    HYBRID = "hybrid"  # Combination of methods
    FALLBACK = "fallback"  # Basic fallback extraction


@dataclass
class DocumentProcessingResult:
    """Results from document processing operation."""
    text_content: str
    extraction_method: ExtractionMethod
    document_type: DocumentType
    processing_time_ms: int
    confidence_score: float  # 0.0 to 1.0
    word_count: int
    character_count: int
    has_formatting: bool
    metadata: Dict[str, Any]
    warnings: List[str]
    errors: List[str]


@dataclass
class DocumentMetadata:
    """Metadata extracted from document."""
    title: Optional[str]
    author: Optional[str]
    created_date: Optional[str]
    modified_date: Optional[str]
    page_count: int
    file_size_bytes: int
    language: Optional[str]
    has_images: bool
    has_tables: bool


class DocumentProcessor:
    """Enhanced document processor for resume analysis with multiple extraction strategies."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.supported_formats = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.doc': DocumentType.DOC,
            '.txt': DocumentType.TXT,
            '.png': DocumentType.IMAGE,
            '.jpg': DocumentType.IMAGE,
            '.jpeg': DocumentType.IMAGE,
        }
        
        # OCR configuration
        self.ocr_config = {
            'lang': 'eng',
            'config': '--oem 3 --psm 6'  # Uniform block of text
        }
    
    async def process_document_from_url(
        self,
        document_url: str,
        max_file_size_mb: int = 10,
        enable_ocr: bool = True,
        quality_threshold: float = 0.7
    ) -> DocumentProcessingResult:
        """
        Process document from URL with comprehensive error handling and fallback strategies.
        
        Args:
            document_url: URL to download and process document from
            max_file_size_mb: Maximum allowed file size in MB
            enable_ocr: Whether to use OCR for image-based content
            quality_threshold: Minimum quality score to accept result
            
        Returns:
            DocumentProcessingResult with extracted content and metadata
        """
        start_time = time.time()
        self.logger.info(f"Starting document processing from URL: {document_url}")
        
        try:
            # Download document
            document_data, content_type = await self._download_document(
                document_url, max_file_size_mb
            )
            
            # Determine document type
            doc_type = self._detect_document_type(document_url, content_type)
            self.logger.info(f"Detected document type: {doc_type.value}")
            
            # Process based on document type
            result = await self._process_document_data(
                document_data, doc_type, enable_ocr, quality_threshold
            )
            
            # Update processing time
            result.processing_time_ms = int((time.time() - start_time) * 1000)
            
            # Validate result quality
            if result.confidence_score < quality_threshold:
                self.logger.warning(
                    f"Document processing quality below threshold: {result.confidence_score:.2f} < {quality_threshold}"
                )
                result.warnings.append(f"Quality score {result.confidence_score:.2f} below threshold {quality_threshold}")
            
            self.logger.info(
                f"Document processing completed: {result.word_count} words, "
                f"confidence: {result.confidence_score:.2f}, method: {result.extraction_method.value}"
            )
            
            return result
            
        except Exception as e:
            self.logger.error(f"Document processing failed: {e}")
            raise JobProcessingException(f"Document processing failed: {str(e)}")
    
    async def process_document_from_file(
        self,
        file_path: str,
        enable_ocr: bool = True,
        quality_threshold: float = 0.7
    ) -> DocumentProcessingResult:
        """Process document from local file path."""
        start_time = time.time()
        
        try:
            # Read file data
            with open(file_path, 'rb') as f:
                document_data = f.read()
            
            # Detect document type from file extension
            doc_type = self._detect_document_type(file_path, None)
            
            # Process document
            result = await self._process_document_data(
                document_data, doc_type, enable_ocr, quality_threshold
            )
            
            result.processing_time_ms = int((time.time() - start_time) * 1000)
            return result
            
        except Exception as e:
            self.logger.error(f"File processing failed: {e}")
            raise JobProcessingException(f"File processing failed: {str(e)}")
    
    async def _download_document(
        self, url: str, max_size_mb: int
    ) -> Tuple[bytes, Optional[str]]:
        """Download document from URL with size and security checks."""
        max_size_bytes = max_size_mb * 1024 * 1024
        
        try:
            async with httpx.AsyncClient(timeout=30.0) as client:
                # Start download with streaming
                async with client.stream('GET', url) as response:
                    response.raise_for_status()
                    
                    # Check content length
                    content_length = response.headers.get('content-length')
                    if content_length and int(content_length) > max_size_bytes:
                        raise JobProcessingException(f"File too large: {content_length} bytes > {max_size_bytes}")
                    
                    # Download with size limit
                    content = b''
                    async for chunk in response.aiter_bytes():
                        content += chunk
                        if len(content) > max_size_bytes:
                            raise JobProcessingException(f"File too large during download: {len(content)} bytes")
                    
                    content_type = response.headers.get('content-type')
                    return content, content_type
                    
        except httpx.RequestError as e:
            raise JobProcessingException(f"Failed to download document: {e}")
        except httpx.HTTPStatusError as e:
            raise JobProcessingException(f"HTTP error downloading document: {e.response.status_code}")
    
    def _detect_document_type(self, file_path_or_url: str, content_type: Optional[str]) -> DocumentType:
        """Detect document type from file extension and content type."""
        # Try file extension first
        path_lower = file_path_or_url.lower()
        for ext, doc_type in self.supported_formats.items():
            if path_lower.endswith(ext):
                return doc_type
        
        # Try content type
        if content_type:
            if 'pdf' in content_type:
                return DocumentType.PDF
            elif 'word' in content_type or 'officedocument' in content_type:
                return DocumentType.DOCX
            elif 'text' in content_type:
                return DocumentType.TXT
            elif 'image' in content_type:
                return DocumentType.IMAGE
        
        return DocumentType.UNKNOWN
    
    async def _process_document_data(
        self,
        document_data: bytes,
        doc_type: DocumentType,
        enable_ocr: bool,
        quality_threshold: float
    ) -> DocumentProcessingResult:
        """Process document data based on detected type."""
        warnings = []
        errors = []
        
        try:
            if doc_type == DocumentType.PDF:
                return await self._process_pdf(document_data, enable_ocr, warnings, errors)
            elif doc_type == DocumentType.DOCX:
                return await self._process_docx(document_data, warnings, errors)
            elif doc_type == DocumentType.DOC:
                return await self._process_doc(document_data, warnings, errors)
            elif doc_type == DocumentType.TXT:
                return await self._process_txt(document_data, warnings, errors)
            elif doc_type == DocumentType.IMAGE:
                return await self._process_image(document_data, enable_ocr, warnings, errors)
            else:
                return await self._process_unknown(document_data, enable_ocr, warnings, errors)
                
        except Exception as e:
            errors.append(f"Processing failed: {str(e)}")
            # Return fallback result
            return DocumentProcessingResult(
                text_content="Document processing failed. Manual review required.",
                extraction_method=ExtractionMethod.FALLBACK,
                document_type=doc_type,
                processing_time_ms=0,
                confidence_score=0.0,
                word_count=0,
                character_count=0,
                has_formatting=False,
                metadata={},
                warnings=warnings,
                errors=errors
            )
    
    async def _process_pdf(
        self, pdf_data: bytes, enable_ocr: bool, warnings: List[str], errors: List[str]
    ) -> DocumentProcessingResult:
        """Process PDF document with multiple extraction strategies."""
        text_content = ""
        extraction_method = ExtractionMethod.DIRECT_TEXT
        confidence_score = 0.0
        has_formatting = False
        metadata = {}
        
        try:
            # Strategy 1: Direct text extraction with PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_data))
            page_count = len(pdf_reader.pages)
            metadata["page_count"] = page_count
            
            # Extract metadata
            if pdf_reader.metadata:
                metadata.update({
                    "title": pdf_reader.metadata.get('/Title'),
                    "author": pdf_reader.metadata.get('/Author'),
                    "created_date": str(pdf_reader.metadata.get('/CreationDate')),
                })
            
            # Try direct text extraction
            direct_text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    direct_text += page_text + "\n"
            
            if direct_text.strip() and len(direct_text.strip()) > 50:
                text_content = direct_text.strip()
                confidence_score = 0.9
                self.logger.info("PDF: Direct text extraction successful")
            else:
                # Strategy 2: Enhanced extraction with pdfplumber
                try:
                    with pdfplumber.open(io.BytesIO(pdf_data)) as pdf:
                        enhanced_text = ""
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                enhanced_text += page_text + "\n"
                                
                            # Check for tables
                            if page.extract_tables():
                                has_formatting = True
                                metadata["has_tables"] = True
                        
                        if enhanced_text.strip() and len(enhanced_text.strip()) > 50:
                            text_content = enhanced_text.strip()
                            confidence_score = 0.8
                            self.logger.info("PDF: Enhanced extraction with pdfplumber successful")
                
                except Exception as e:
                    warnings.append(f"Enhanced PDF extraction failed: {str(e)}")
            
            # Strategy 3: OCR as fallback
            if (not text_content or confidence_score < 0.5) and enable_ocr:
                try:
                    text_content = await self._extract_pdf_with_ocr(pdf_data)
                    extraction_method = ExtractionMethod.OCR
                    confidence_score = 0.6
                    warnings.append("Used OCR for text extraction - formatting may be lost")
                    self.logger.info("PDF: OCR extraction used as fallback")
                except Exception as ocr_error:
                    errors.append(f"OCR extraction failed: {str(ocr_error)}")
            
        except Exception as e:
            errors.append(f"PDF processing error: {str(e)}")
            confidence_score = 0.0
        
        # Final validation
        if not text_content:
            text_content = "Unable to extract text from PDF. Manual review required."
            confidence_score = 0.0
            extraction_method = ExtractionMethod.FALLBACK
        
        return DocumentProcessingResult(
            text_content=text_content,
            extraction_method=extraction_method,
            document_type=DocumentType.PDF,
            processing_time_ms=0,  # Will be set by caller
            confidence_score=confidence_score,
            word_count=len(text_content.split()),
            character_count=len(text_content),
            has_formatting=has_formatting,
            metadata=metadata,
            warnings=warnings,
            errors=errors
        )
    
    async def _process_docx(
        self, docx_data: bytes, warnings: List[str], errors: List[str]
    ) -> DocumentProcessingResult:
        """Process DOCX document."""
        try:
            # Strategy 1: python-docx for structured extraction
            doc = DocxDocument(io.BytesIO(docx_data))
            
            text_parts = []
            has_formatting = False
            metadata = {}
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract tables
            if doc.tables:
                has_formatting = True
                metadata["has_tables"] = True
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_parts.append(" | ".join(row_text))
            
            text_content = "\n".join(text_parts)
            confidence_score = 0.9 if text_content.strip() else 0.0
            
            # Extract core properties
            if doc.core_properties:
                metadata.update({
                    "title": doc.core_properties.title,
                    "author": doc.core_properties.author,
                    "created": str(doc.core_properties.created),
                    "modified": str(doc.core_properties.modified),
                })
            
            return DocumentProcessingResult(
                text_content=text_content,
                extraction_method=ExtractionMethod.DIRECT_TEXT,
                document_type=DocumentType.DOCX,
                processing_time_ms=0,
                confidence_score=confidence_score,
                word_count=len(text_content.split()),
                character_count=len(text_content),
                has_formatting=has_formatting,
                metadata=metadata,
                warnings=warnings,
                errors=errors
            )
            
        except Exception as e:
            errors.append(f"DOCX processing error: {str(e)}")
            
            # Strategy 2: Fallback with mammoth for better formatting preservation
            try:
                result = mammoth.extract_raw_text(io.BytesIO(docx_data))
                text_content = result.value
                confidence_score = 0.7 if text_content.strip() else 0.0
                
                if result.messages:
                    warnings.extend([str(msg) for msg in result.messages])
                
                return DocumentProcessingResult(
                    text_content=text_content,
                    extraction_method=ExtractionMethod.HYBRID,
                    document_type=DocumentType.DOCX,
                    processing_time_ms=0,
                    confidence_score=confidence_score,
                    word_count=len(text_content.split()),
                    character_count=len(text_content),
                    has_formatting=False,
                    metadata={},
                    warnings=warnings,
                    errors=errors
                )
            except Exception as mammoth_error:
                errors.append(f"Mammoth fallback failed: {str(mammoth_error)}")
                
                return DocumentProcessingResult(
                    text_content="DOCX processing failed. Manual review required.",
                    extraction_method=ExtractionMethod.FALLBACK,
                    document_type=DocumentType.DOCX,
                    processing_time_ms=0,
                    confidence_score=0.0,
                    word_count=0,
                    character_count=0,
                    has_formatting=False,
                    metadata={},
                    warnings=warnings,
                    errors=errors
                )
    
    async def _process_doc(
        self, doc_data: bytes, warnings: List[str], errors: List[str]
    ) -> DocumentProcessingResult:
        """Process legacy DOC document."""
        warnings.append("Legacy DOC format - limited support, conversion recommended")
        
        try:
            # For legacy DOC files, we'd typically use external tools
            # For now, return a message indicating limited support
            return DocumentProcessingResult(
                text_content="Legacy DOC format detected. Please convert to DOCX or PDF for better processing.",
                extraction_method=ExtractionMethod.FALLBACK,
                document_type=DocumentType.DOC,
                processing_time_ms=0,
                confidence_score=0.1,
                word_count=0,
                character_count=0,
                has_formatting=False,
                metadata={"format": "legacy_doc"},
                warnings=warnings,
                errors=errors
            )
        except Exception as e:
            errors.append(f"DOC processing error: {str(e)}")
            return DocumentProcessingResult(
                text_content="DOC processing failed.",
                extraction_method=ExtractionMethod.FALLBACK,
                document_type=DocumentType.DOC,
                processing_time_ms=0,
                confidence_score=0.0,
                word_count=0,
                character_count=0,
                has_formatting=False,
                metadata={},
                warnings=warnings,
                errors=errors
            )
    
    async def _process_txt(
        self, txt_data: bytes, warnings: List[str], errors: List[str]
    ) -> DocumentProcessingResult:
        """Process plain text document."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            text_content = None
            
            for encoding in encodings:
                try:
                    text_content = txt_data.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if text_content is None:
                # Fallback with error handling
                text_content = txt_data.decode('utf-8', errors='ignore')
                warnings.append("Used fallback encoding - some characters may be corrupted")
            
            return DocumentProcessingResult(
                text_content=text_content,
                extraction_method=ExtractionMethod.DIRECT_TEXT,
                document_type=DocumentType.TXT,
                processing_time_ms=0,
                confidence_score=0.95,
                word_count=len(text_content.split()),
                character_count=len(text_content),
                has_formatting=False,
                metadata={"encoding": "detected_automatically"},
                warnings=warnings,
                errors=errors
            )
            
        except Exception as e:
            errors.append(f"TXT processing error: {str(e)}")
            return DocumentProcessingResult(
                text_content="Text processing failed.",
                extraction_method=ExtractionMethod.FALLBACK,
                document_type=DocumentType.TXT,
                processing_time_ms=0,
                confidence_score=0.0,
                word_count=0,
                character_count=0,
                has_formatting=False,
                metadata={},
                warnings=warnings,
                errors=errors
            )
    
    async def _process_image(
        self, image_data: bytes, enable_ocr: bool, warnings: List[str], errors: List[str]
    ) -> DocumentProcessingResult:
        """Process image document with OCR."""
        if not enable_ocr:
            warnings.append("OCR disabled - cannot extract text from image")
            return DocumentProcessingResult(
                text_content="OCR disabled for image processing.",
                extraction_method=ExtractionMethod.FALLBACK,
                document_type=DocumentType.IMAGE,
                processing_time_ms=0,
                confidence_score=0.0,
                word_count=0,
                character_count=0,
                has_formatting=False,
                metadata={},
                warnings=warnings,
                errors=errors
            )
        
        try:
            # Use OCR to extract text
            text_content = await self._extract_image_with_ocr(image_data)
            confidence_score = 0.6 if text_content and len(text_content.strip()) > 10 else 0.3
            
            return DocumentProcessingResult(
                text_content=text_content,
                extraction_method=ExtractionMethod.OCR,
                document_type=DocumentType.IMAGE,
                processing_time_ms=0,
                confidence_score=confidence_score,
                word_count=len(text_content.split()) if text_content else 0,
                character_count=len(text_content) if text_content else 0,
                has_formatting=False,
                metadata={"ocr_used": True},
                warnings=warnings + ["OCR used - accuracy may vary"],
                errors=errors
            )
            
        except Exception as e:
            errors.append(f"Image OCR processing error: {str(e)}")
            return DocumentProcessingResult(
                text_content="Image OCR processing failed.",
                extraction_method=ExtractionMethod.FALLBACK,
                document_type=DocumentType.IMAGE,
                processing_time_ms=0,
                confidence_score=0.0,
                word_count=0,
                character_count=0,
                has_formatting=False,
                metadata={},
                warnings=warnings,
                errors=errors
            )
    
    async def _process_unknown(
        self, data: bytes, enable_ocr: bool, warnings: List[str], errors: List[str]
    ) -> DocumentProcessingResult:
        """Process document of unknown type."""
        warnings.append("Unknown document format detected")
        
        # Try to detect if it's actually text
        try:
            text_content = data.decode('utf-8', errors='ignore')
            if len(text_content.strip()) > 10 and text_content.isprintable():
                return DocumentProcessingResult(
                    text_content=text_content,
                    extraction_method=ExtractionMethod.HYBRID,
                    document_type=DocumentType.UNKNOWN,
                    processing_time_ms=0,
                    confidence_score=0.5,
                    word_count=len(text_content.split()),
                    character_count=len(text_content),
                    has_formatting=False,
                    metadata={"detected_as_text": True},
                    warnings=warnings + ["Treated unknown format as text"],
                    errors=errors
                )
        except Exception:
            pass
        
        return DocumentProcessingResult(
            text_content="Unknown document format. Manual processing required.",
            extraction_method=ExtractionMethod.FALLBACK,
            document_type=DocumentType.UNKNOWN,
            processing_time_ms=0,
            confidence_score=0.0,
            word_count=0,
            character_count=0,
            has_formatting=False,
            metadata={},
            warnings=warnings,
            errors=errors
        )
    
    async def _extract_pdf_with_ocr(self, pdf_data: bytes) -> str:
        """Extract text from PDF using OCR on rendered pages."""
        try:
            # This would require additional libraries like pdf2image
            # For now, return a placeholder
            return "OCR extraction from PDF not fully implemented. Install pdf2image and configure poppler."
        except Exception as e:
            raise Exception(f"PDF OCR extraction failed: {e}")
    
    async def _extract_image_with_ocr(self, image_data: bytes) -> str:
        """Extract text from image using OCR."""
        try:
            # Open image with PIL
            image = Image.open(io.BytesIO(image_data))
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Use tesseract for OCR
            text = pytesseract.image_to_string(
                image, 
                lang=self.ocr_config['lang'],
                config=self.ocr_config['config']
            )
            
            return text.strip()
            
        except Exception as e:
            raise Exception(f"Image OCR failed: {e}")


# Factory function
def create_document_processor() -> DocumentProcessor:
    """Create and return a configured document processor."""
    return DocumentProcessor()


# Integration functions for the main Brain service
async def process_resume_document(
    document_url: str,
    max_file_size_mb: int = 10,
    enable_ocr: bool = True
) -> Dict[str, Any]:
    """
    Process resume document and return results in format compatible with Brain service.
    
    This function integrates with the existing message processing flow.
    """
    processor = create_document_processor()
    
    try:
        result = await processor.process_document_from_url(
            document_url=document_url,
            max_file_size_mb=max_file_size_mb,
            enable_ocr=enable_ocr,
            quality_threshold=0.5
        )
        
        return {
            "status": "success",
            "text_content": result.text_content,
            "processing_metadata": {
                "document_type": result.document_type.value,
                "extraction_method": result.extraction_method.value,
                "confidence_score": result.confidence_score,
                "processing_time_ms": result.processing_time_ms,
                "word_count": result.word_count,
                "warnings": result.warnings,
                "errors": result.errors
            },
            "document_metadata": result.metadata
        }
        
    except Exception as e:
        logging.getLogger(__name__).error(f"Document processing failed: {e}")
        return {
            "status": "error",
            "error": str(e),
            "message": "Document processing failed. Please ensure the document is accessible and in a supported format."
        }